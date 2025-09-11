from fastapi import FastAPI, UploadFile, File, Path, Query, BackgroundTasks, Form, HTTPException
from typing import List, Dict, Tuple
import io
import cv2
import numpy as np
import pdfplumber
import pytesseract
from PIL import Image, ImageSequence
import pydicom
import spacy
from docx import Document
import pandas as pd
import json
from pdf2image import convert_from_path
import tempfile
import os
from fastapi.responses import JSONResponse, FileResponse
import re
import httpx
from pydantic import BaseModel
import hashlib
from datetime import datetime, timezone
from sqlalchemy import create_engine, Column, Integer, String, DateTime
from sqlalchemy.orm import declarative_base, sessionmaker
from twilio.rest import Client
from dotenv import load_dotenv
import asyncio
from collections import Counter
import uuid
import fitz
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from fastapi.middleware.cors import CORSMiddleware
from spacy.pipeline import EntityRuler

# Init FastAPI
app = FastAPI(title="MedVault Multi-Modal Medical Document Processor")

# Allow frontend to talk to backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Storing progress for batch proccessing 
progress_store = {}

# Load NLP model for PII detection
nlp = spacy.load("en_core_web_md")

# Create EntityRuler
ruler = nlp.add_pipe("entity_ruler", before="ner")

patterns = [
    # Medical conditions (simplified list, can extend)
    {"label": "CONDITION", "pattern": [{"LOWER": "hypertension"}]},
    {"label": "CONDITION", "pattern": [{"LOWER": "diabetes"}]},
    {"label": "CONDITION", "pattern": [{"LOWER": "asthma"}]},
    {"label": "CONDITION", "pattern": [{"LOWER": "fever"}]}, 

    # Insurance identifiers
    {"label": "POLICY", "pattern": [{"LOWER": "policy"}, {"LOWER": "no"}, {"IS_DIGIT": True}]},
    {"label": "CLAIM", "pattern": [{"LOWER": "claim"}, {"IS_DIGIT": True}]},
    {"label": "ACCOUNT", "pattern": [{"LOWER": "account"}, {"IS_DIGIT": True}]},

    # Legal identifiers
    {"label": "CASE", "pattern": [{"LOWER": "case"}, {"IS_DIGIT": True}]},
    {"label": "LAW", "pattern": [{"LOWER": "section"}, {"IS_DIGIT": True}]},
    {"label": "COURT", "pattern": [{"LOWER": "high"}, {"LOWER": "court"}]}
]

ruler.add_patterns(patterns)

# Define entity groups
PII_ENTITIES = {"PERSON", "GPE", "ORG", "DATE", "TIME", "LOC", "NORP"}
PHI_ENTITIES = PII_ENTITIES.union({"CONDITION"})
INSURANCE_ENTITIES = PII_ENTITIES.union({"POLICY", "CLAIM", "ACCOUNT"})
LEGAL_ENTITIES = PII_ENTITIES.union({"LAW", "CASE", "COURT"})

# Map modes → what to redact
MODE_ENTITY_MAP = {
    "research": PII_ENTITIES,          # anonymize patients but keep conditions
    "patient": PHI_ENTITIES,           # full PHI redaction
    "insurance": INSURANCE_ENTITIES,   # redact PII + insurance info
    "legal": LEGAL_ENTITIES            # redact PII + legal IDs
}

# Load environment variables and twirlio credentials
load_dotenv()
TWILIO_ACCOUNT_SID = os.getenv("TWILIO_ACCOUNT_SID")
TWILIO_AUTH_TOKEN = os.getenv("TWILIO_AUTH_TOKEN")
TWILIO_PHONE = os.getenv("TWILIO_PHONE")
ALERT_PHONE = os.getenv("ALERT_PHONE")

twilio_client = Client(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)

# Models
class Document(BaseModel):
    id: str
    content: str

class Action(BaseModel):
    action: str
    user: str

#  Database  Setup
DATABASE_URL = "sqlite:///./medvault_audit.db"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class AuditLog(Base):
    __tablename__ = "audit_logs"
    id = Column(Integer, primary_key=True, index=True)
    doc_id = Column(String)
    action = Column(String)
    user = Column(String)
    timestamp = Column(DateTime, default=datetime.now(timezone.utc))
    fingerprint = Column(String)

Base.metadata.create_all(bind=engine)

# Blockchain Setup

class Block:
    def __init__(self, index, timestamp, data, previous_hash):
        self.index = index
        self.timestamp = timestamp
        self.data = data
        self.previous_hash = previous_hash
        self.hash = self.compute_hash()
    def compute_hash(self):
        return hashlib.sha256(f"{self.index}{self.timestamp}{self.data}{self.previous_hash}".encode()).hexdigest()

blockchain = []
def create_genesis_block():
    return Block(0, datetime.now(timezone.utc).isoformat(), "Genesis Block", "0")
if not blockchain:
    blockchain.append(create_genesis_block())

async def add_block_async(data):
    prev_block = blockchain[-1]
    new_block = Block(len(blockchain), datetime.now(timezone.utc).isoformat(), data, prev_block.hash)
    blockchain.append(new_block)
    return new_block

# HIPAA Compliance (18 identifiers)
HIPAA_IDENTIFIERS = {
    "names": r"\b([A-Z][a-z]+ [A-Z][a-z]+)\b",
    "ssn": r"\b\d{3}-\d{2}-\d{4}\b",
    "phone": r"\b\d{3}-\d{3}-\d{4}\b",
    "email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
    "address": r"\d{1,5} [A-Za-z0-9\s]+ (Street|St|Avenue|Ave|Rd|Road|Blvd|Lane|Ln)\b",
    "dates": r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
    "medical_record_number": r"\bMRN\d+\b",
    "health_plan_number": r"\bHP\d+\b",
    "account_numbers": r"\bAC\d+\b",
    "certificate_numbers": r"\bCERT\d+\b",
    "license_numbers": r"\bLIC\d+\b",
    "vehicle_ids": r"\bVIN[A-Z0-9]+\b",
    "device_ids": r"\bDEV[A-Z0-9]+\b",
    "web_urls": r"https?://[^\s]+",
    "ip_addresses": r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b",
    "biometric_identifiers": r"\bFINGERPRINT\b|\bIRIS\b",
    "full_face_photos": r"\bPHOTO\b",
    "any_other_unique_id": r"\bUID\d+\b"
}

def check_hipaa_compliance(text: str) -> List[str]:
    found = []
    for key, pattern in HIPAA_IDENTIFIERS.items():
        if re.search(pattern, text):
            found.append(key)
    return found

# ---------- Utility: Redact Text ----------
def redact_text(text: str, mode: str = "research") -> str:
    doc = nlp(text)
    entities_to_redact = MODE_ENTITY_MAP.get(mode, set())

    redacted = text
    for ent in doc.ents:
        if ent.label_ in entities_to_redact:
            # Use regex to avoid partial replacements messing up
            redacted = redacted.replace(ent.text, "[REDACTED]")
    
    return redacted
# ---------- OCR Function ----------
def extract_text(file_path: str):
    ext = file_path.split(".")[-1].lower()
    text = ""

    if ext in ["jpg", "jpeg", "png", "tiff"]:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)

    elif ext == "pdf":
        images = convert_from_path(file_path)
        for img in images:
            text += pytesseract.image_to_string(img) + "\n"

    return text

# ---------- NER Function ----------
def detect_entities(text: str):
    doc = nlp(text)
    entities = []
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "GPE", "ORG", "DATE", "CARDINAL"]:
            entities.append({"text": ent.text, "label": ent.label_})
    return entities

# ---------- Computer Vision ----------
def detect_sensitive_regions(file_path: str):
    img = cv2.imread(file_path)

    results = {"faces": 0, "signatures": 0}

    # Face detection
    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + "haarcascade_frontalface_default.xml")
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    faces = face_cascade.detectMultiScale(gray, 1.2, 5)
    results["faces"] = len(faces)

    # Simple signature/stamp detection (by contours, very basic)
    gray_blur = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(gray_blur, 50, 150)
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    for c in contours:
        x, y, w, h = cv2.boundingRect(c)
        if w > 100 and h < 100:  # heuristic for signature shape
            results["signatures"] += 1

    return results

# ----------Function for different Privacy modes ----------
def apply_privacy_mode(text: str, entities: List[Dict], mode: str) -> str:
    """
    Redacts sensitive info depending on the privacy mode.
    """

    redacted_text = text

    if mode == "patient":
        # Keep patient’s own info but redact others
        for ent in entities:
            if ent["label"] in ["PERSON", "NAME"] and not ent["is_patient"]:
                redacted_text = redacted_text.replace(ent["text"], "[REDACTED]")
    
    elif mode == "research":
        # Full de-identification
        for ent in entities:
            redacted_text = redacted_text.replace(ent["text"], "[REDACTED]")
        # Date shifting (simple example)
        redacted_text = re.sub(r"\d{4}-\d{2}-\d{2}", "[SHIFTED_DATE]", redacted_text)
        # Location anonymization
        redacted_text = re.sub(r"\b(?:New York|California|Bangalore)\b", "[LOCATION]", redacted_text)

    elif mode == "insurance":
        # Keep only treatment dates + claim info
        filtered = []
        for ent in entities:
            if ent["label"] in ["DATE", "TREATMENT", "CLAIM"]:
                filtered.append(f"{ent['label']}: {ent['text']}")
        return "\n".join(filtered)

    elif mode == "legal":
        # Strong redaction for compliance
        for ent in entities:
            redacted_text = redacted_text.replace(ent["text"], "[LEGAL_REDACTED]")

    return redacted_text

# --- Different medical document category recognition ---
DOC_CATEGORIES = {
    "discharge_summary": [
        r"\badmission date\b", r"\bdischarge date\b", r"\bhospital course\b",
        r"\bchief complaint\b", r"\bdisposition\b", r"\bdischarge medications?\b",
        r"\bprimary diagnosis\b", r"\battending physician\b"
    ],
    "lab_report": [
        r"\bspecimen\b", r"\bcollected\b", r"\breceived\b", r"\banalyte\b",
        r"\bresult\b", r"\breference range\b", r"\bunits?\b", r"\btest code\b",
        r"\bOBX\|", r"\bOBR\|"
    ],
    "radiology_report": [
        r"\bimpression\b", r"\bfindings\b", r"\btechnique\b", r"\bcomparison\b",
        r"\bmodality\b", r"\bCT\b", r"\bMRI\b", r"\bx[- ]?ray\b", r"\bultrasound\b"
    ],
    "operative_note": [
        r"\bpre[- ]?operative diagnosis\b", r"\bpost[- ]?operative diagnosis\b",
        r"\bprocedure\b", r"\bsurgeon\b", r"\bassistant\b", r"\banesthesia\b",
        r"\bestimated blood loss\b", r"\bspecimens?\b"
    ],
    "prescription": [
        r"\bRx\b", r"\bSig\b", r"\bDisp(?:ense)?\b", r"\bRefills?\b",
        r"\bNPI\b", r"\bDEA\b", r"\bq\d+h\b", r"\bmg\b", r"\btablet\b", r"\bcapsule\b"
    ],
    "progress_note": [
        r"\bSubjective:\b", r"\bObjective:\b", r"\bAssessment:\b", r"\bPlan:\b",  # SOAP
        r"\bHPI\b", r"\bROS\b", r"\bPE\b", r"\bfollow[- ]?up\b"
    ],
    "referral_letter": [
        r"\bDear Dr\b", r"\breferr?al\b", r"\bconsult(?:ation)?\b", r"\battn\b",
        r"\bI am referring\b", r"\bfor evaluation of\b"
    ],
    "insurance_claim": [
        r"\bclaim number\b", r"\bCMS[- ]?1500\b", r"\bHCFA\b", r"\bEOB\b",
        r"\bICD[- ]?10\b", r"\bCPT\b", r"\bpayer\b", r"\bdeductible\b",
        r"\bcoinsurance\b", r"\bpolicy number\b", r"\bmember id\b"
    ],
    "consent_form": [
        r"\bconsent\b", r"\bI hereby\b", r"\bvoluntarily\b", r"\brisk[s]?\b",
        r"\bbenefit[s]?\b", r"\bwitness\b", r"\bpatient signature\b",
        r"\bguardian\b", r"\bauthorize\b", r"\bdisclosure\b"
    ],
    "billing_invoice": [
        r"\binvoice\b", r"\bamount due\b", r"\bbalance\b", r"\bdate of service\b",
        r"\bcharges?\b", r"\bpayments?\b", r"\badjustments?\b"
    ],
}

# Optional heading/structure boosters
HEADINGS = [
    r"^\s*impression\s*:", r"^\s*findings\s*:", r"^\s*technique\s*:", r"^\s*assessment\s*:", r"^\s*plan\s*:",
    r"^\s*chief complaint\s*:", r"^\s*hospital course\s*:", r"^\s*disposition\s*:", r"^\s*procedure\s*:",
    r"^\s*diagnosis\s*:"
]

def _score_category(text: str, patterns: List[str]) -> float:
    score = 0.0
    for pat in patterns:
        matches = re.findall(pat, text, flags=re.IGNORECASE | re.MULTILINE)
        if matches:
            # Base presence points + frequency factor
            score += 2.0 + 0.5 * len(matches)
    return score

def _heading_bonus(text: str) -> float:
    bonus = 0.0
    for pat in HEADINGS:
        if re.search(pat, text, flags=re.IGNORECASE | re.MULTILINE):
            bonus += 0.75
    return bonus

def classify_document(text: str) -> Dict:
    """
    Returns:
      {
        'label': 'lab_report',
        'confidence': 0.88,
        'scores': {'lab_report': 6.5, 'radiology_report': 2.0, ...},
        'evidence': ['matched: "reference range"', 'matched: "OBX|"', ...]
      }
    """
    if not text or not text.strip():
        return {"label": "unknown", "confidence": 0.0, "scores": {}, "evidence": []}

    scores = {}
    evidence: List[str] = []
    norm = text

    # Score all categories
    for cat, pats in DOC_CATEGORIES.items():
        s = _score_category(norm, pats)
        scores[cat] = s

    # Mild global heading bonus to all (helps clinical formats)
    hb = _heading_bonus(norm)
    for cat in scores:
        scores[cat] += hb * 0.25

    # Track evidence by showing top-matching regex tokens
    for cat, pats in DOC_CATEGORIES.items():
        for pat in pats:
            if re.search(pat, norm, flags=re.IGNORECASE | re.MULTILINE):
                # add a short explanation once per pattern
                literal = re.sub(r"\\b|\?:|\(|\)|\[|\]|\||\+|\*|\^|\$|\\", "", pat)
                evidence.append(f'{cat}: matched "{literal[:32]}{"..." if len(literal)>32 else ""}"')

    # Normalize to probabilities
    total = sum(scores.values()) or 1.0
    probs = {k: (v / total) for k, v in scores.items()}
    label = max(probs.items(), key=lambda kv: kv[1])[0]
    confidence = round(probs[label], 4)

    # If everything is super low, call it unknown
    if scores[label] < 1.5:  # tune threshold
        label, confidence = "unknown", 0.0

    # Keep only top 6 evidence strings to keep response tidy
    return {
        "label": label,
        "confidence": confidence,
        "scores": {k: round(v, 3) for k, v in scores.items()},
        "evidence": evidence[:6]
    }

# ---------- Send SMS ----------
def send_sms(message: str):
    twilio_client.messages.create(
        body=message,
        from_=TWILIO_PHONE,
        to=ALERT_PHONE
    )

# Audit and Compliance For all types of files 
async def add_audit_entry_async(doc_id: str, action: str, user: str):
    timestamp = datetime.now(timezone.utc)
    fingerprint = hashlib.sha256(f"{doc_id}{action}{timestamp}".encode()).hexdigest()
    db = SessionLocal()
    entry = AuditLog(
        doc_id=doc_id,
        action=action,
        user=user,
        timestamp=timestamp,
        fingerprint=fingerprint
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)
    db.close()
    return entry

async def audit_file(file_content: str, filename: str, user: str, background_tasks: BackgroundTasks):
    # HIPAA compliance check
    violations = check_hipaa_compliance(file_content)
    risk = "high" if violations else "low"

    # DB audit log
    entry = await add_audit_entry_async(filename, "audit_check", user)

    # Blockchain logging
    blockchain_entry = await add_block_async({
        "doc_id": filename,
        "action": "file_processed",
        "user": user,
        "violations": violations
    })

    # SMS alert if violations found
    if violations:
        message = f"HIPAA violations detected in file {filename}: {violations}"
        background_tasks.add_task(send_sms, message)

    return {
        "violations": violations,
        "risk": risk,
        "audit_log": {
            "doc_id": entry.doc_id,
            "action": entry.action,
            "user": entry.user,
            "timestamp": entry.timestamp,
            "fingerprint": entry.fingerprint
        },
        "blockchain_hash": blockchain_entry.hash
    }

# Utility to save and return download path
def save_redacted_file(content: bytes, ext: str) -> str:
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uuid.uuid4()}{ext}")
    with open(temp_file.name, "wb") as f:
        f.write(content)
    return temp_file.name

# ---------- PDF Processing ----------
@app.post("/process/pdf")
async def process_pdf(
    file: UploadFile = File(...),
    user: str = "admin",
    background_tasks: BackgroundTasks = None,
    privacy_mode: str = "research"
):
    contents = await file.read()
    pages_text = []

    # Try extracting text page by page
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

    # Fallback: OCR if no text layer found in any page
    if not pages_text:
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page in pdf.pages:
                image = page.to_image(resolution=300).original
                text = pytesseract.image_to_string(image)
                pages_text.append(text)

    # Redact each page individually
    redacted_pages = [redact_text(page, mode=privacy_mode) for page in pages_text]

    # Join all pages for classification + auditing
    full_text = "\n".join(pages_text)
    audit_info = await audit_file(full_text, file.filename, user, background_tasks)
    classification = classify_document(full_text)

    # --- Create redacted PDF (preserving page structure) ---
    redacted_dir = os.path.join(tempfile.gettempdir(), "redacted")
    os.makedirs(redacted_dir, exist_ok=True)

    output_path = os.path.join(redacted_dir, file.filename)
    packet = io.BytesIO()
    c = canvas.Canvas(packet, pagesize=letter)

    for page_text in redacted_pages:
        text_object = c.beginText(40, 750)  # margins
        for line in page_text.split("\n"):
            text_object.textLine(line)
        c.drawText(text_object)
        c.showPage()  # new page for next
    c.save()

    with open(output_path, "wb") as f:
        f.write(packet.getvalue())

    return {
        "original_pages": [p[:500] for p in pages_text],      # first 500 chars per page
        "redacted_pages": [r[:500] for r in redacted_pages],  # redacted preview per page
        "compliance": audit_info,
        "privacy_mode": privacy_mode,
        "classification": classification,
        "page_count": len(pages_text),
        "download_url": f"/download/{file.filename}" # endpoint to fetch file
    }

# ---------- Image Processing (JPEG, PNG, TIFF) ----------
@app.post("/process/image")
async def process_image(
    file: UploadFile = File(...),
    user: str = "admin",
    background_tasks: BackgroundTasks = None,
    privacy_mode: str = "research"
):
    contents = await file.read()
    np_img = np.frombuffer(contents, np.uint8)

    # Try to decode as a single-page image
    img = cv2.imdecode(np_img, cv2.IMREAD_COLOR)
    texts = []

    if img is not None:
        # OCR with bounding boxes
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        text = " ".join(data["text"])
        texts.append(text)

        # (Optional) mask all detected words → if you want to blackout everything
        for i, word in enumerate(data["text"]):
            if word.strip():
                (x, y, w, h) = (
                    data["left"][i],
                    data["top"][i],
                    data["width"][i],
                    data["height"][i],
                )
                cv2.rectangle(img, (x, y), (x + w, y + h), (0, 0, 0), -1)

        # Save redacted image
        redacted_dir = os.path.join(tempfile.gettempdir(), "redacted")
        os.makedirs(redacted_dir, exist_ok=True)

        redacted_filename = f"redacted_{uuid.uuid4()}.png"
        redacted_path = os.path.join(redacted_dir, redacted_filename)
        cv2.imwrite(redacted_path, img)

    else:
        # Multi-page TIFF handling
        pil_img = Image.open(io.BytesIO(contents))
        frame_texts = []
        for frame in range(0, getattr(pil_img, "n_frames", 1)):
            pil_img.seek(frame)
            frame_text = pytesseract.image_to_string(pil_img)
            frame_texts.append(frame_text)
        texts.extend(frame_texts)

        # Save TIFF without any bounding box redaction
        redacted_filename = f"redacted_{uuid.uuid4()}.tiff"
        redacted_path = os.path.join("/tmp", redacted_filename)
        pil_img.save(redacted_path)

    full_text = "\n".join(texts)

    # Use your existing text redaction + classification
    redacted_text = redact_text(full_text, mode=privacy_mode)
    classification = classify_document(full_text)
    audit_info = await audit_file(full_text, file.filename, user, background_tasks)

    return {
        "original": full_text[:500],
        "redacted": redacted_text[:500],
        "compliance": audit_info,
        "privacy_mode": privacy_mode,
        "classification": classification,
        "pages": len(texts),
        "download_url": f"/download/{redacted_filename}"  # 👈 allows download
    }

# ---------- DICOM Medical Scan Processing ----------
@app.post("/process/dicom")
async def process_dicom(
    files: List[UploadFile] = File(...),
    user: str = "admin",
    background_tasks: BackgroundTasks = None,
    privacy_mode: str = "research"
):
    results = []
    output_files = []

    # Define DICOM tags to redact by mode
    MODE_DICOM_TAGS = {
        "research": ["PatientName", "PatientID"],  
        "patient": ["PatientName", "PatientID", "PatientBirthDate", "PatientSex"],
        "insurance": ["PatientName", "PatientID", "PatientBirthDate"],
        "legal": ["PatientName", "PatientID", "PatientBirthDate", "PatientSex"]
    }

    # Default tags if mode not found
    tags_to_redact = MODE_DICOM_TAGS.get(privacy_mode, ["PatientName", "PatientID"])

    redacted_dir = os.path.join(tempfile.gettempdir(), "redacted")
    os.makedirs(redacted_dir, exist_ok=True)

    for file in files:
        contents = await file.read()
        ds = pydicom.dcmread(io.BytesIO(contents))

        # Extract metadata (before redaction)
        metadata = {elem.keyword: str(elem.value) for elem in ds if elem.keyword}

        # Apply redaction
        for tag in tags_to_redact:
            if tag in ds:
                ds.data_element(tag).value = "REDACTED"

        # Save redacted DICOM
        output_path = os.path.join(redacted_dir, file.filename)
        ds.save_as(output_path)
        output_files.append(output_path)

        # Collect redacted text for audit/classification
        pii_text = " ".join([str(ds.get(tag, "")) for tag in tags_to_redact])
        classification = classify_document(pii_text)
        audit_info = await audit_file(pii_text, file.filename, user, background_tasks)

        results.append({
            "filename": file.filename,
            "metadata": metadata,
            "message": f"Redacted {tags_to_redact}",
            "compliance": audit_info,
            "privacy_mode": privacy_mode,
            "classification": classification,
            "download_url": f"/download/{file.filename}"
        })

    return {"results": results}

# ---------- Word Documents (Clinical Notes, Emails) ----------
@app.post("/process/word")
async def process_word(
    file: UploadFile = File(...),
    user: str = "admin",
    background_tasks: BackgroundTasks = None,
    privacy_mode: str = "research"
):
    contents = await file.read()
    doc = Document(io.BytesIO(contents))

    # Extract text page-wise (simulate multi-page by section breaks or paragraphs)
    pages = []
    current_page = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current_page.append(text)
        # Treat "Page Break" as separator
        if "pageBreakBefore" in para._element.xml:  
            pages.append("\n".join(current_page))
            current_page = []
    if current_page:
        pages.append("\n".join(current_page))

    # If no explicit page breaks, treat the whole document as one "page"
    if not pages:
        pages = ["\n".join([p.text for p in doc.paragraphs])]

    # Process each page
    results = []
    for i, page_text in enumerate(pages, start=1):
        redacted = redact_text(page_text, mode=privacy_mode)
        classification = classify_document(page_text)
        audit_info = await audit_file(page_text, f"{file.filename}_page_{i}", user, background_tasks)
        
        results.append({
            "page": i,
            "original": page_text[:500],
            "redacted": redacted[:500],
            "compliance": audit_info,
            "privacy_mode": privacy_mode,
            "classification": classification
        })

    return {
        "filename": file.filename,
        "total_pages": len(pages),
        "results": results
    }

# ---------- Excel/CSV (Lab Results) ----------
@app.post("/process/sheet")
async def process_sheet(
    file: UploadFile = File(...),
    user: str = "admin",
    background_tasks: BackgroundTasks = None,
    privacy_mode: str = "research"
):
    contents = await file.read()

    text_blocks = []
    try:
        # Try reading as Excel with multiple sheets
        xls = pd.ExcelFile(io.BytesIO(contents))
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            text_blocks.append(f"--- Sheet: {sheet_name} ---\n{df.to_string(index=False)}")
    except Exception:
        # If not Excel, fallback to CSV
        try:
            df = pd.read_csv(io.BytesIO(contents))
            text_blocks.append(f"--- CSV File ---\n{df.to_string(index=False)}")
        except Exception as e:
            return {"error": f"Unable to parse file: {str(e)}"}

    full_text = "\n\n".join(text_blocks)

    # Apply redaction and classification
    redacted = redact_text(full_text, mode=privacy_mode)
    classification = classify_document(full_text)
    audit_info = await audit_file(full_text, file.filename, user, background_tasks)

    return {
        "original": full_text[:1000],   # send preview only
        "redacted": redacted[:1000],
        "compliance": audit_info,
        "privacy_mode": privacy_mode,
        "classification": classification,
        "sheets": len(text_blocks)      # number of sheets processed
    }

# ---------- HL7/FHIR Structured JSON ----------
@app.post("/process/hl7")
async def process_hl7(
    file: UploadFile = File(...),
    user: str = "admin",
    background_tasks: BackgroundTasks = None,
    privacy_mode: str = "research"
):
    contents = await file.read()
    data = json.loads(contents)

    def recursive_redact(d):
        if isinstance(d, dict):
            return {k: recursive_redact(v) for k, v in d.items()}
        elif isinstance(d, list):
            return [recursive_redact(i) for i in d]
        elif isinstance(d, str):
            return redact_text(d, mode=privacy_mode)
        else:
            return d

    redacted = recursive_redact(data)
    classification = classify_document(json.dumps(data))
    audit_info = await audit_file(json.dumps(data), file.filename, user, background_tasks)

    # Convert dicts → string for preview
    data_str = json.dumps(data, indent=2)
    redacted_str = json.dumps(redacted, indent=2)

    return {
        "original": data_str[:500],      # ✅ safe preview
        "redacted": redacted_str[:500],  # ✅ safe preview
        "compliance": audit_info,
        "privacy_mode": privacy_mode,
        "classification": classification
    }

# ---------- Upload any number and type of documents ----------
@app.post("/upload")
async def upload_files(
    files: list[UploadFile] = File(...),   # Accept multiple files
    privacy_mode: str = Form("research"),
    user: str = Form("admin"),
    background_tasks: BackgroundTasks = None
):
    batch_id = str(uuid.uuid4())
    progress_store[batch_id] = {"total": len(files), "processed": 0, "results": []}

    results = []

    for file in files:
        suffix = os.path.splitext(file.filename)[-1].lower()

        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        try:
            await file.seek(0)

            if suffix == ".pdf":
                result = await process_pdf(file=file, privacy_mode=privacy_mode, user=user, background_tasks=background_tasks)
            elif suffix in [".docx", ".doc"]:
                result = await process_word(file=file, privacy_mode=privacy_mode, user=user, background_tasks=background_tasks)
            elif suffix in [".jpg", ".jpeg", ".png", ".tiff"]:
                result = await process_image(file=file, privacy_mode=privacy_mode, user=user, background_tasks=background_tasks)
            elif suffix == ".dcm":
                result = await process_dicom(files=files, privacy_mode=privacy_mode, user=user, background_tasks=background_tasks)
            elif suffix in [".xlsx", ".xls", ".csv"]:
                result = await process_sheet(file=file, privacy_mode=privacy_mode, user=user, background_tasks=background_tasks)
            elif suffix in [".json", ".hl7"]:
                result = await process_hl7(file=file, privacy_mode=privacy_mode, user=user, background_tasks=background_tasks)
            else:
                result = {"error": f"Unsupported file type: {suffix}"}

            results.append({file.filename: result})

            # Update progress
            progress_store[batch_id]["processed"] += 1
            progress_store[batch_id]["results"].append({file.filename: result})

        finally:
            os.remove(tmp_path)

    return {"batch_id": batch_id, "results": results}


@app.get("/upload/progress/{batch_id}")
async def get_batch_progress(batch_id: str):
    progress = progress_store.get(batch_id)
    if not progress:
        return JSONResponse({"error": "Invalid batch_id"}, status_code=404)

    return {
        "batch_id": batch_id,
        "processed": progress["processed"],
        "total": progress["total"],
        "results": progress["results"]
    }

# ---------------- EMR (FHIR Patients) ----------------
@app.get("/emr/patients", summary="Get patients (EMR - HAPI FHIR Sandbox)")
async def get_patients():
    async with httpx.AsyncClient() as client:
        r = await client.get("https://hapi.fhir.org/baseR4/Patient")
        return r.json()

# ---------------- Labs (Observations) ----------------
@app.get("/labs/observations", summary="Get laboratory observations")
async def get_labs(
    category: str = Query("laboratory", description="Category of observation"),
    count: int = Query(5, alias="_count", description="Number of records")
):
    async with httpx.AsyncClient() as client:
        r = await client.get(
            "https://hapi.fhir.org/baseR4/Observation",
            params={"category": category, "_count": count}
        )
        return r.json()

# ---------------- PACS (DICOM Images) ----------------
@app.get("/pacs/studies", summary="Get imaging studies (mock)")
async def get_studies():
    # Mock data simulating what Orthanc DICOMweb would return
    studies = [
        {
            "StudyDate": "20250831",
            "StudyTime": "101530",
            "Modality": "CT",
            "PatientName": "John Doe",
            "StudyURL": "https://mock-pacs/studies/1.2.840.113619.2.5.1762583153.1234.1592403528.467",
        },
        {
            "StudyDate": "20250830",
            "StudyTime": "143210",
            "Modality": "MR",
            "PatientName": "Jane Smith",
            "StudyURL": "https://mock-pacs/studies/2.25.123456789012345678901234567890123456",
        }
    ]
    
    return {"count": len(studies), "studies": studies}


@app.get("/pacs/studies/{studyId}/series", summary="Get series for a study (mock)")
async def get_study_series(
    studyId: str = Path(..., description="Study Instance UID")
):
    # Mock series data for the given studyId
    series = [
        {
            "SeriesInstanceUID": "1.2.840.113619.2.55.3.604688319.783.1592403528.112",
            "SeriesDescription": "Head CT without contrast",
            "Modality": "CT",
            "BodyPartExamined": "BRAIN",
        },
        {
            "SeriesInstanceUID": "1.2.840.113619.2.55.3.604688319.784.1592403528.113",
            "SeriesDescription": "Head CT with contrast",
            "Modality": "CT",
            "BodyPartExamined": "BRAIN",
        }
    ]
    
    return {"studyId": studyId, "count": len(series), "series": series}

# ---------------- Insurance / Claims ----------------
@app.get("/insurance/claims", summary="Get insurance claims")
async def get_claims():
    return [
        {"id": "claim-001", "patient": "Rohan Sharma", "amount": 15200, "status": "Pending"},
        {"id": "claim-002", "patient": "Aditi Verma", "amount": 19800, "status": "Approved"},
        {"id": "claim-003", "patient": "Arjun Mehta", "amount": 22400, "status": "Rejected"},
        {"id": "claim-004", "patient": "Priya Nair", "amount": 17600, "status": "Pending"},
        {"id": "claim-005", "patient": "Karan Gupta", "amount": 30500, "status": "Approved"},
        {"id": "claim-006", "patient": "Neha Iyer", "amount": 18750, "status": "Pending"},
        {"id": "claim-007", "patient": "Siddharth Reddy", "amount": 24200, "status": "Approved"},
        {"id": "claim-008", "patient": "Ananya Mukherjee", "amount": 26800, "status": "Rejected"},
        {"id": "claim-009", "patient": "Vikram Singh", "amount": 28900, "status": "Pending"},
        {"id": "claim-010", "patient": "Meera Pillai", "amount": 31500, "status": "Approved"},
    ]

# ---------- logs + blockchain + risk check in one request ----------
@app.post("/audit/process")
async def process_document(doc: Document, action: Action, background_tasks: BackgroundTasks):
    # 1. HIPAA compliance check
    violations = check_hipaa_compliance(doc.content)
    risk = "high" if violations else "low"

    # 2. Async DB audit log
    entry = await add_audit_entry_async(doc.id, action.action, action.user)

    # 3. Async blockchain logging
    blockchain_entry = await add_block_async({
        "doc_id": doc.id,
        "action": action.action,
        "user": action.user,
        "violations": violations
    })

    # 4. Send SMS alert if violations found
    if violations:
        message = f"HIPAA violations detected in doc {doc.id}: {violations}"
        background_tasks.add_task(send_sms, message)

    return {
        "hipaa_compliant": risk == "low",
        "violations": violations,
        "audit_log": {
            "doc_id": entry.doc_id,
            "action": entry.action,
            "user": entry.user,
            "timestamp": entry.timestamp,
            "fingerprint": entry.fingerprint
        },
        "blockchain_hash": blockchain_entry.hash,
        "risk": risk
    }

# ---------- Document Classification ----------
@app.post("/classify/text")
async def classify_text_endpoint(payload: dict):
    """
    payload = { "text": "..." }
    """
    text = payload.get("text", "")
    result = classify_document(text)
    return result

@app.post("/classify/file")
async def classify_file_endpoint(file: UploadFile = File(...)):
    """
    Tries to extract text depending on file suffix and then classifies.
    Uses a very small extractor set, mirroring your existing logic style
    (PDF/text layer -> OCR, Word, CSV/Excel, DICOM metadata, images).
    """
    suffix = os.path.splitext(file.filename)[-1].lower()
    content = await file.read()

    extracted_text = ""

    try:
        if suffix == ".pdf":
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for p in pdf.pages:
                    t = p.extract_text() or ""
                    extracted_text += t + "\n"
            if not extracted_text.strip():
                # OCR fallback
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for p in pdf.pages:
                        img = p.to_image(resolution=300).original
                        extracted_text += pytesseract.image_to_string(img) + "\n"

        elif suffix in [".jpg", ".jpeg", ".png", ".tiff"]:
            np_img = np.frombuffer(content, np.uint8)
            img = cv2.imdecode(np_img, cv2.IMREAD_COLOR)
            extracted_text = pytesseract.image_to_string(img)

        elif suffix in [".docx", ".doc"]:
            doc = Document(io.BytesIO(content))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])

        elif suffix in [".xlsx", ".xls", ".csv"]:
            # Try excel first; fallback to csv
            try:
                df = pd.read_excel(io.BytesIO(content))
            except:
                df = pd.read_csv(io.BytesIO(content))
            extracted_text = df.to_string()

        elif suffix == ".dcm":
            ds = pydicom.dcmread(io.BytesIO(content))
            # Use metadata string for classification
            fields = []
            for elem in ds:
                if getattr(elem, "keyword", None):
                    v = str(elem.value)
                    if isinstance(v, bytes):
                        v = v.decode("utf-8", "ignore")
                    fields.append(f"{elem.keyword}: {v}")
            extracted_text = "\n".join(fields)

        elif suffix in [".json", ".hl7"]:
            try:
                data = json.loads(content.decode("utf-8", "ignore"))
            except Exception:
                # plain text HL7 pipe format as fallback
                data = content.decode("utf-8", "ignore")
            extracted_text = json.dumps(data, ensure_ascii=False) if isinstance(data, (dict, list)) else str(data)

        else:
            return JSONResponse({"error": f"Unsupported file type: {suffix}"}, status_code=400)

    finally:
        # Reset file pointer for any further use elsewhere
        await file.seek(0)

    result = classify_document(extracted_text)
    # Include a tiny preview so you can verify quickly in Swagger
    result["preview"] = extracted_text[:500]
    result["filename"] = file.filename
    return result

# ---------- Blockchain Viewing Endpoints ----------
@app.get("/blockchain", summary="Get full blockchain")
async def get_blockchain():
    """
    Returns the current blockchain entries for transparency.
    Each entry includes doc_id, action, user, violations, timestamp, and hash.
    """
    return {"chain": blockchain}

# ---------- Verify Blockchain Integrity ----------
@app.get("/blockchain/verify")
async def verify_blockchain():
    """
    Verifies the blockchain integrity.
    Returns true if the chain is valid (no tampering).
    """
    for i in range(1, len(blockchain)):
        current = blockchain[i]
        prev = blockchain[i - 1]

        # 1. Check hash integrity
        if current.hash != current.compute_hash():
            return {"valid": False, "error": f"Invalid hash at block {i}"}

        # 2. Check chain linkage
        if current.previous_hash != prev.hash:
            return {"valid": False, "error": f"Broken chain link at block {i}"}

    return {"valid": True, "message": "Blockchain integrity verified"}

# ---------- Download Redacted File ----------
@app.get("/download/{filename}")
async def download_file(filename: str):
    redacted_dir = os.path.join(tempfile.gettempdir(), "redacted")
    file_path = os.path.join(redacted_dir, filename)

    if not os.path.exists(file_path):
        return {"error": f"File {filename} not found"}

    return FileResponse(file_path, filename=filename, media_type="application/pdf")
# ---------- Root Endpoint ----------
@app.get("/")
def root():
    return {"message": "MedVault Multi-Modal Processor is running"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)