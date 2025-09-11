from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import pandas as pd
import json
import pydicom
from pydicom.dataset import Dataset, FileDataset
import datetime
import os

# Create output directory
out_dir = "medvault_test_files"
os.makedirs(out_dir, exist_ok=True)

# 1. Sample PDF
pdf_path = os.path.join(out_dir, "prescription.pdf")
c = canvas.Canvas(pdf_path, pagesize=letter)
c.drawString(100, 750, "Patient: John Doe")
c.drawString(100, 730, "Date: 2025-08-31")
c.drawString(100, 710, "Diagnosis: Fever")
c.save()

# 2. Sample Word Document
docx_path = os.path.join(out_dir, "clinical_note.docx")
doc = Document()
doc.add_paragraph("Patient admitted: David Brown")
doc.add_paragraph("Date: 2025-08-30")
doc.add_paragraph("Treatment: Antibiotics")
doc.save(docx_path)

# 3. Sample CSV
csv_path = os.path.join(out_dir, "lab_results.csv")
df = pd.DataFrame({
    "Patient": ["Jane Doe"],
    "Date": ["2025-08-30"],
    "Result": ["Positive"]
})
df.to_csv(csv_path, index=False)

# 4. Sample JSON (HL7/FHIR)
json_path = os.path.join(out_dir, "hl7.json")
hl7_data = {
    "resourceType": "Patient",
    "name": "Michael Scott",
    "birthDate": "1985-03-15",
    "address": "Scranton, USA"
}
with open(json_path, "w") as f:
    json.dump(hl7_data, f, indent=4)

# 5. Sample DICOM
dicom_path = os.path.join(out_dir, "scan.dcm")
file_meta = Dataset()
file_meta.MediaStorageSOPClassUID = pydicom.uid.ExplicitVRLittleEndian
file_meta.MediaStorageSOPInstanceUID = pydicom.uid.generate_uid()
file_meta.ImplementationClassUID = pydicom.uid.generate_uid()

ds = FileDataset(dicom_path, {}, file_meta=file_meta, preamble=b"\0" * 128)
ds.PatientName = "Robert Smith"
ds.PatientID = "123456"
ds.PatientBirthDate = "19700101"
ds.PatientSex = "M"
ds.StudyDate = datetime.date.today().strftime("%Y%m%d")
ds.Modality = "CT"
ds.is_little_endian = True
ds.is_implicit_VR = False
ds.save_as(dicom_path)

print("✅ Test files generated in:", out_dir)
